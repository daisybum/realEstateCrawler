#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX Parser for Real Estate Crawler
Handles extraction of text, tables, and images from DOCX files
"""

import os
import logging
import tempfile
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple

import pytesseract
from PIL import Image
import docx
import numpy as np

from src.config import Config


class DOCXParser:
    """Parser for DOCX files"""
    
    def __init__(self, config=None):
        """Initialize the DOCX parser"""
        self.config = config or Config.get_instance()
        self.logger = logging.getLogger(__name__)
        
    def parse_bytes(self, file_obj) -> Dict[str, Any]:
        """
        Parse DOCX from BytesIO object
        
        Args:
            file_obj: BytesIO object containing DOCX data
            
        Returns:
            Dictionary with parsed content
        """
        try:
            result = {
                "content": "",
                "metadata": {},
                "tables": [],
                "images": []
            }
            
            # Process DOCX
            doc = docx.Document(file_obj)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            result["content"] = "\n\n".join(paragraphs)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell_idx, cell in enumerate(row.cells):
                        row_data.append(cell.text)
                    table_data.append(row_data)
                
                if table_data:
                    # Convert to dictionary format
                    if len(table_data) > 1:  # Has header and data
                        headers = [str(h).strip() for h in table_data[0]]
                        dict_data = []
                        for row in table_data[1:]:
                            row_dict = {}
                            for i, cell in enumerate(row):
                                if i < len(headers):
                                    row_dict[headers[i]] = str(cell).strip()
                            dict_data.append(row_dict)
                            
                        result["tables"].append({
                            'index': table_idx,
                            'headers': headers,
                            'data': dict_data
                        })
                    else:
                        # No header, just raw data
                        result["tables"].append({
                            'index': table_idx,
                            'data': table_data
                        })
            
            # Extract images using zipfile approach
            try:
                from zipfile import ZipFile
                file_obj.seek(0)  # Reset file pointer
                
                with ZipFile(file_obj) as docx_zip:
                    # Get list of all image files in the document
                    image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                    
                    for img_idx, img_path in enumerate(image_files):
                        try:
                            # Extract image data
                            image_bytes = docx_zip.read(img_path)
                            
                            # Save to temporary file for OCR
                            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_file:
                                tmp_file.write(image_bytes)
                                image_path = tmp_file.name
                            
                            # Perform OCR
                            ocr_text = self._perform_ocr(image_path)
                            
                            # Add to images list
                            result["images"].append({
                                'index': img_idx,
                                'filename': img_path.split('/')[-1],
                                'ocr_text': ocr_text
                            })
                            
                            # Clean up
                            os.unlink(image_path)
                        except Exception as e:
                            self.logger.warning(f"Error extracting image {img_path}: {e}")
            except Exception as e:
                self.logger.warning(f"Error extracting images from DOCX: {e}")
            
            # Extract metadata
            try:
                core_properties = doc.core_properties
                result["metadata"] = {
                    'title': core_properties.title or '',
                    'author': core_properties.author or '',
                    'subject': core_properties.subject or '',
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables)
                }
            except Exception as e:
                self.logger.warning(f"Error extracting metadata: {e}")
            
            return result
            
        except Exception as e:
            self.logger.error(f"Error parsing DOCX bytes: {e}")
            return {"error": str(e), "content": f"DOCX 파일 처리 오류: {e}"}
    
    def _perform_ocr(self, image_path: str) -> str:
        """
        Perform OCR on an image file
        
        Args:
            image_path: Path to the image file
            
        Returns:
            Extracted text from the image
        """
        try:
            # Open image with PIL
            img = Image.open(image_path)
            
            # Perform OCR with pytesseract
            text = pytesseract.image_to_string(img, lang='kor+eng')
            return text.strip()
        except Exception as e:
            self.logger.warning(f"OCR failed: {e}")
            return ""
    
    def extract_content(self, docx_path: str) -> Dict[str, Any]:
        """
        Extract content from a DOCX file
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dictionary containing extracted content
        """
        if not os.path.exists(docx_path):
            self.logger.error(f"DOCX file not found: {docx_path}")
            return {"error": f"DOCX file not found: {docx_path}"}
        
        try:
            result = {
                "text": "",
                "paragraphs": [],
                "tables": [],
                "images": [],
                "metadata": {}
            }
            
            # Extract text, tables, and metadata
            text, paragraphs, tables, metadata = self._extract_docx_content(docx_path)
            result["text"] = text
            result["paragraphs"] = paragraphs
            result["tables"] = tables
            result["metadata"] = metadata
            
            # Extract images (if any)
            images = self._extract_images(docx_path)
            result["images"] = images
            
            return result
        
        except Exception as e:
            self.logger.error(f"Error extracting content from DOCX: {e}", exc_info=True)
            return {"error": f"Error extracting content from DOCX: {str(e)}"}
    
    def _extract_docx_content(self, docx_path: str) -> Tuple[str, List[Dict[str, Any]], List[Dict[str, Any]], Dict[str, Any]]:
        """
        Extract content from DOCX file
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Tuple of (text, paragraphs, tables, metadata)
        """
        text = ""
        paragraphs = []
        tables = []
        metadata = {}
        
        try:
            # Open the document
            doc = docx.Document(docx_path)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision or 0,
                "category": core_props.category or "",
                "content_status": core_props.content_status or "",
                "language": core_props.language or "",
            }
            
            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    para_text = para.text.strip()
                    text += para_text + "\n\n"
                    
                    paragraphs.append({
                        "index": i,
                        "text": para_text,
                        "style": para.style.name if para.style else "Normal"
                    })
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = []
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    
                    table_data.append(row_data)
                
                tables.append({
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.rows[0].cells) if table.rows else 0,
                    "data": table_data
                })
                
                # Add table representation to text
                text += f"\n[Table {i+1} with {len(table.rows)} rows]\n\n"
            
            return text, paragraphs, tables, metadata
        
        except Exception as e:
            self.logger.error(f"Error extracting DOCX content: {e}", exc_info=True)
            return "", [], [], {}
    
    def _extract_images(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        Extract images from DOCX file
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            List of images as dictionaries with OCR text
        """
        images_data = []
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Note: python-docx doesn't provide direct access to images
            # This is a placeholder for actual implementation
            # In a real implementation, you might:
            # 1. Extract the DOCX as a ZIP file
            # 2. Access the word/media directory
            # 3. Process each image file
            
            self.logger.warning("DOCX image extraction not fully implemented")
            
            # For now, just return an empty list
            return images_data
            
        except Exception as e:
            self.logger.error(f"Error extracting images from DOCX: {e}", exc_info=True)
            
        return images_data
